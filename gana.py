from docx import Document
from docx.shared import Inches
import pyttsx3

pyttsx3.speak('hello')


document = Document() 

document.add_picture(
    'gan.jpg',
    width=Inches(1.0)
)

name = input('what is your name ? ')
phone_number = input('what is your phone number?')
email = input('what is your email? ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)

document.add_heading('About me')
about_me = input('Tell about yourself? ')
document.add_paragraph(about_me)

document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('enter company ')
start_date = input('From date ')
to_date = input('To date ')

p.add_run(company + ' ').bold = True
p.add_run(start_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'describe your experience at ' +  company
)

p.add_run(experience_details)
document.add_heading('Skills')
while True:
    has_more_experiences = input(
        'do you have more experiences ? yes or no')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('enter company ')
        start_date = input('From date ')
        to_date = input('To date ')

        p.add_run(company + ' ').bold = True
        p.add_run(start_date + '-' + to_date + '\n').italic = True

        experience_details = input(
           'describe your experience at ' +  company)

        p.add_run(experience_details)
    
    else:
        break

while True:
    has_skills = input(
        'do you have any extra skills? yes or no'
    )
    if has_skills == 'yes':
        p = document.add_paragraph()

        skill = input('enter your skill: ')
        level = input('eneter at which level you know that skill?')
        certification = input('enter yes or no to this')

        p.add_run(skill + ' ').italic = True
        p.add_run(level + ' ').bold = True
        p.add_run(certification + ' ' + '\n').bold = True

    else:
        break


       
    
 

document.save('cv.docx')